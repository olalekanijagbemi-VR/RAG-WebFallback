"""
RAG+WebFallback - Main Application (FINAL - MERGED)
Multi-Agent RAG System with Web Fallback
Apple Liquid Glass UI - Dark Bars - 3D Italic Title - Big Bold Input
"""

import os
import re
import time
import logging
import math
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple
from collections import Counter

import numpy as np
import pandas as pd
import streamlit as st
from dotenv import load_dotenv

# --- Vector / keyword search (Pure Python BM25) ---
from sklearn.neighbors import NearestNeighbors

# --- Chunking (pure-python, no torch) ---
from langchain_text_splitters import RecursiveCharacterTextSplitter

# --- Document parsers ---
import pypdf
import docx

# --- Gemini SDK ---
from google import genai
from google.genai import types

# --- Web fallback ---
from ddgs import DDGS

# --- Env / secrets ---
load_dotenv()

# ============================================================
# CONFIG
# ============================================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("RAG_Web_Fallback")

APP_TITLE = "Multi-Agent RAG with Web Fallback"

CHUNK_SIZE_TOKENS = 500
CHUNK_OVERLAP_TOKENS = 50
CHARS_PER_TOKEN = 4

EMBED_MODEL = "models/gemini-embedding-001"
GEN_MODEL = "models/gemini-3.5-flash"

SEMANTIC_WEIGHT = 0.6
KEYWORD_WEIGHT = 0.4

TOP_K_DOCS = 5
WEB_MAX_RESULTS = 5

TIME_SENSITIVE_KEYWORDS = [
    "today", "current", "currently", "latest", "news", "weather", "now",
    "this week", "this month", "this year", "recent", "recently", "live",
    "score", "stock price", "price of", "update", "breaking",
    "right now", "up to date", "real-time", "real time", "forecast",
    "who won", "what happened", "trending",
    "where is", "what is", "who is", "capital of", "population of",
    "which country", "largest", "smallest", "tallest", "longest",
    "located in", "located at", "geography", "city", "country",
    "when was", "how to", "how does", "meaning of",
]

SYSTEM_PROMPT = """You are a precise, careful research assistant.

Rules:
1. Only use information contained in the SOURCES provided below.
2. Cite every factual claim using the matching [Source X] notation.
3. If the sources do not contain enough information to answer confidently,
   say so plainly instead of guessing.
4. If sources conflict, point out the discrepancy rather than picking one silently.
5. Be concise, well-organized, and use markdown formatting where useful.
"""


# ============================================================
# APPLE LIQUID GLASS UI - DARK BARS - 3D ITALIC TITLE
# ============================================================
st.set_page_config(
    page_title="RAG+WebFallback",
    page_icon="⚙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    /* ============================================================
       BACKGROUND: High Resolution Reflective Metal at 75% intensity
       ============================================================ */
    
    .stApp {
        background-color: #4a4a4a;
        
        /* High Resolution Metal Texture */
        background-image: 
            repeating-linear-gradient(90deg, 
                rgba(0,0,0,0.02) 0px, 
                rgba(255,255,255,0.02) 0.5px, 
                transparent 1px, 
                transparent 3px,
                rgba(0,0,0,0.015) 3px,
                rgba(255,255,255,0.015) 3.5px,
                transparent 4px,
                transparent 6px
            ),
            repeating-linear-gradient(90deg, 
                rgba(0,0,0,0.03) 0px, 
                transparent 2px, 
                rgba(255,255,255,0.02) 4px, 
                transparent 6px
            ),
            linear-gradient(90deg, 
                #3a3a3a 0%, 
                #6a6a6a 25%, 
                #b0b0b0 40%, 
                #cccccc 50%, 
                #b0b0b0 60%, 
                #6a6a6a 75%, 
                #3a3a3a 100%
            );
            
        background-blend-mode: overlay, overlay, normal;
        background-size: cover, cover, cover;
        background-attachment: fixed;
    }

    /* ============================================================
       TOP SEAM - Double Width (12px) - Dark
       ============================================================ */
    
    .stApp::before {
        content: "";
        position: fixed;
        top: 10%;
        left: 0;
        width: 100%;
        height: 12px;
        z-index: 9999;
        pointer-events: none;
        background: 
            linear-gradient(to bottom, 
                rgba(0,0,0,0.8) 0px, 
                rgba(0,0,0,0.95) 3px, 
                rgba(60,50,40,0.4) 4px,
                rgba(180,170,160,0.3) 5px,
                rgba(255,255,255,0.6) 6px, 
                rgba(255,255,255,0.3) 7px,
                rgba(180,170,160,0.15) 8px,
                transparent 12px
            );
    }

    /* ============================================================
       BOTTOM SEAM - Double Width (12px) - Dark
       ============================================================ */
    
    .stApp::after {
        content: "";
        position: fixed;
        bottom: 10%;
        left: 0;
        width: 100%;
        height: 12px;
        z-index: 9999;
        pointer-events: none;
        background: 
            linear-gradient(to bottom, 
                transparent 0px,
                rgba(180,170,160,0.15) 4px,
                rgba(255,255,255,0.3) 5px,
                rgba(255,255,255,0.6) 6px, 
                rgba(180,170,160,0.3) 7px,
                rgba(60,50,40,0.4) 8px,
                rgba(0,0,0,0.95) 9px, 
                rgba(0,0,0,0.8) 12px
            );
    }

    /* ============================================================
       TOP BAR - DARK BLACK (NO WHITE, NO BLUE)
       ============================================================ */
    
    header[data-testid="stHeader"] {
        background: rgba(0,0,0,0.95) !important;
        backdrop-filter: none !important;
        border-bottom: 1px solid rgba(255,255,255,0.05) !important;
        box-shadow: 0px 2px 20px rgba(0,0,0,0.8) !important;
        height: 48px !important;
        min-height: 48px !important;
    }
    
    .stApp > header {
        background: rgba(0,0,0,0.95) !important;
    }
    
    .st-emotion-cache-1r6slb0 {
        background: rgba(0,0,0,0.95) !important;
    }

    /* ============================================================
       HIDE FOOTER ONLY - KEEP CHAT INPUT
       ============================================================ */

    footer {
        display: none !important;
        visibility: hidden !important;
        height: 0 !important;
    }
    
    #MainMenu {
        display: none !important;
        visibility: hidden !important;
    }
    
    .st-emotion-cache-1r6slb0 {
        display: none !important;
    }

    /* ============================================================
       CUSTOM BLACK BOTTOM BAR (THIN, MATCHES TOP)
       ============================================================ */

    body::after {
        content: "";
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        height: 48px;
        z-index: 999999;
        pointer-events: none;
        background: linear-gradient(180deg, 
            #1a1a1a 0%, 
            #2a2a2a 25%, 
            #1a1a1a 50%, 
            #111111 75%, 
            #0a0a0a 100%
        );
        background-image: 
            radial-gradient(ellipse at 5% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 15% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 25% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 35% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 45% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 55% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 65% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 75% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 85% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            radial-gradient(ellipse at 95% 50%, rgba(50,50,50,0.4) 2px, rgba(30,30,30,0.2) 3px, transparent 4px),
            repeating-linear-gradient(90deg, 
                rgba(0,0,0,0.05) 0px, 
                rgba(255,255,255,0.02) 1px, 
                transparent 2px, 
                transparent 6px
            );
        background-blend-mode: overlay, overlay, normal;
        background-size: auto, cover;
        background-repeat: repeat-x, no-repeat;
        border-top: 2px solid rgba(0,0,0,0.5);
        box-shadow: 0px -2px 10px rgba(0,0,0,0.6);
    }

    /* ============================================================
       MAIN TITLE - 3D Effect, Shadow, Slightly Italic
       Gear Icon stays STRAIGHT (not italic)
       ============================================================ */
    
    .main-header {
        font-size: 3.2rem;
        font-weight: 800;
        color: #979799 !important;
        letter-spacing: 2px;
        margin-bottom: 0.2rem;
        padding: 0.5rem 0;
        display: flex;
        align-items: center;
        gap: 0.3rem;
    }
    
    .main-header .gear-icon {
        font-style: normal !important;
        font-weight: 400;
        display: inline-block;
        transform: none !important;
        text-shadow: 
            0px 2px 4px rgba(0,0,0,0.2),
            0px 4px 12px rgba(0,0,0,0.15) !important;
    }
    
    .main-header .title-text {
        font-style: italic !important;
        text-shadow: 
            0px 1px 0px rgba(0,0,0,0.1),
            0px 2px 0px rgba(0,0,0,0.15),
            0px 3px 0px rgba(0,0,0,0.2),
            0px 4px 0px rgba(0,0,0,0.25),
            0px 5px 0px rgba(0,0,0,0.3),
            0px 6px 0px rgba(0,0,0,0.35),
            0px 8px 12px rgba(0,0,0,0.5),
            0px 12px 24px rgba(0,0,0,0.3) !important;
        transform: skewX(-3deg);
        display: inline-block;
    }
    
    .sub-header {
        font-size: 1.3rem;
        color: #FF0000 !important;
        text-shadow: 0px 2px 8px rgba(0,0,0,0.3) !important;
        margin-bottom: 2rem;
        opacity: 1;
        font-weight: 500;
    }

    /* ============================================================
       RESPONSE TEXT - Very Dark and BOLD
       ============================================================ */
    
    .stChatMessage div, 
    .stChatMessage p, 
    .stChatMessage span,
    .stChatMessage .stMarkdown {
        color: #0a0a0a !important;
        font-weight: 700 !important;
        text-shadow: 
            0px 1px 2px rgba(255, 255, 255, 0.08),
            0px 2px 8px rgba(255, 255, 255, 0.03) !important;
        letter-spacing: 0.01em;
    }
    
    .stChatMessage .stMarkdown p {
        color: #0a0a0a !important;
        font-weight: 700 !important;
    }

    /* ============================================================
       APPLE LIQUID GLASS - Response Layer
       ============================================================ */
    
    .stChatMessage[data-testid="stChatMessage"]:has(div[data-testid="stChatMessageContent"]:nth-child(1)) {
        background: rgba(255, 255, 255, 0.15) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(255, 255, 255, 0.2) !important;
        border-radius: 1.2rem !important;
        padding: 1rem 1.5rem !important;
        margin: 0.5rem 0 !important;
        box-shadow: 
            0px 4px 24px rgba(0, 0, 0, 0.08),
            0px 1px 0px rgba(255, 255, 255, 0.3) inset !important;
    }
    
    .stChatMessage[data-testid="stChatMessage"]:has(div[data-testid="stChatMessageContent"]:nth-child(2)) {
        background: rgba(220, 220, 235, 0.12) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(255, 255, 255, 0.15) !important;
        border-radius: 1.2rem !important;
        padding: 1rem 1.5rem !important;
        margin: 0.5rem 0 !important;
        box-shadow: 
            0px 4px 24px rgba(0, 0, 0, 0.06),
            0px 1px 0px rgba(255, 255, 255, 0.2) inset !important;
    }

    /* ============================================================
       CHAT INPUT - APPLE LIQUID GLASS STYLE (VISIBLE)
       ============================================================ */
    
    .stChatInput {
        position: fixed !important;
        bottom: 60px !important;
        left: 50% !important;
        transform: translateX(-50%) !important;
        width: 80% !important;
        max-width: 800px !important;
        z-index: 999999 !important;
        padding: 0 !important;
    }
    
    .stChatInput > div {
        background: rgba(255, 255, 255, 0.15) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(255, 255, 255, 0.25) !important;
        border-radius: 1.2rem !important;
        box-shadow: 
            0px 4px 30px rgba(0, 0, 0, 0.15),
            0px 1px 0px rgba(255, 255, 255, 0.3) inset !important;
        transition: all 0.3s ease !important;
        padding: 0.25rem !important;
    }
    
    .stChatInput > div:focus-within {
        background: rgba(255, 255, 255, 0.22) !important;
        border-color: rgba(255, 255, 255, 0.4) !important;
        box-shadow: 
            0px 4px 40px rgba(0, 0, 0, 0.2),
            0px 1px 0px rgba(255, 255, 255, 0.4) inset !important;
    }
    
    .stChatInput input {
        color: #0a0a0a !important;
        background: transparent !important;
        font-weight: 700 !important;
        font-size: 1.2rem !important;
        text-shadow: 
            0px 1px 2px rgba(255, 255, 255, 0.1) !important;
        letter-spacing: 0.02em;
        padding: 0.75rem 1.2rem !important;
        height: 56px !important;
    }
    
    .stChatInput input::placeholder {
        color: rgba(0, 0, 0, 0.35) !important;
        opacity: 0.8;
        font-weight: 400;
        font-size: 1rem !important;
        text-shadow: none !important;
    }

    /* ============================================================
       SIDEBAR - KEEP ORIGINAL (White text, glass effect)
       ============================================================ */
    
    section[data-testid="stSidebar"] {
        background: rgba(10, 10, 20, 0.7) !important;
        backdrop-filter: blur(8px);
        -webkit-backdrop-filter: blur(8px);
        border-right: 1px solid rgba(255, 255, 255, 0.06);
        z-index: 99999 !important;
    }
    
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] .stMarkdown,
    section[data-testid="stSidebar"] label {
        color: #ffffff !important;
        text-shadow: 0px 2px 8px rgba(0,0,0,0.9) !important;
    }

    /* ============================================================
       SOURCE BOX - Apple Liquid Glass style
       ============================================================ */
    
    .source-box {
        background: rgba(255, 255, 255, 0.08) !important;
        backdrop-filter: blur(12px) !important;
        -webkit-backdrop-filter: blur(12px) !important;
        padding: 0.6rem 1rem;
        border-radius: 0.75rem;
        margin: 0.3rem 0;
        font-size: 0.9rem;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        color: #0a0a0a !important;
        text-shadow: none !important;
        box-shadow: 0px 2px 12px rgba(0,0,0,0.06);
    }
    
    .source-box b {
        color: #0a0a0a !important;
    }
    
    .source-box span {
        color: #0a0a0a !important;
        text-shadow: none !important;
    }

    /* ============================================================
       CONFIDENCE COLORS - Deep Bold Colors
       ============================================================ */
    
    .confidence-high {
        color: #0a6e1a !important;
        font-weight: 700;
        text-shadow: none !important;
    }
    
    .confidence-medium {
        color: #8a6d00 !important;
        font-weight: 700;
        text-shadow: none !important;
    }
    
    .confidence-low {
        color: #8a1a1a !important;
        font-weight: 700;
        text-shadow: none !important;
    }

    /* ============================================================
       BUTTONS - Dark Matte
       ============================================================ */
    
    .stButton > button {
        background: linear-gradient(180deg, #2a2a2a, #1a1a1a) !important;
        border: 1px solid rgba(255,255,255,0.06) !important;
        color: #ffffff !important;
        text-shadow: 0px 1px 4px rgba(0,0,0,0.8) !important;
        border-radius: 0.75rem !important;
        transition: all 0.2s ease !important;
        box-shadow: 0px 2px 12px rgba(0,0,0,0.4) !important;
    }
    
    .stButton > button:hover {
        background: linear-gradient(180deg, #3a3a3a, #2a2a2a) !important;
        box-shadow: 0px 4px 20px rgba(0,0,0,0.5) !important;
        transform: translateY(-1px);
        border-color: rgba(255,255,255,0.12) !important;
    }

    /* ============================================================
       EXPANDER - Apple Liquid Glass style
       ============================================================ */
    
    .streamlit-expanderHeader {
        background: rgba(255, 255, 255, 0.06) !important;
        backdrop-filter: blur(12px) !important;
        -webkit-backdrop-filter: blur(12px) !important;
        border-radius: 0.75rem !important;
        border: 1px solid rgba(255, 255, 255, 0.08) !important;
        color: #0a0a0a !important;
        text-shadow: none !important;
    }
    
    .streamlit-expanderContent {
        background: rgba(255, 255, 255, 0.04) !important;
        backdrop-filter: blur(8px) !important;
        -webkit-backdrop-filter: blur(8px) !important;
        border-radius: 0 0 0.75rem 0.75rem !important;
        border: 1px solid rgba(255, 255, 255, 0.04) !important;
        border-top: none !important;
    }

    /* ============================================================
       FILE UPLOADER - Apple Liquid Glass style
       ============================================================ */
    
    .stFileUploader > div {
        background: rgba(255, 255, 255, 0.05) !important;
        backdrop-filter: blur(8px) !important;
        -webkit-backdrop-filter: blur(8px) !important;
        border: 1px dashed rgba(255, 255, 255, 0.1) !important;
        border-radius: 0.75rem !important;
        color: #0a0a0a !important;
    }

    /* ============================================================
       METRIC CARDS - Apple Liquid Glass style
       ============================================================ */
    
    .stMetric {
        background: rgba(255, 255, 255, 0.05) !important;
        backdrop-filter: blur(8px) !important;
        -webkit-backdrop-filter: blur(8px) !important;
        border-radius: 0.75rem !important;
        padding: 0.5rem 1rem !important;
        border: 1px solid rgba(255, 255, 255, 0.06) !important;
        box-shadow: 0px 2px 12px rgba(0,0,0,0.06);
    }
    
    .stMetric label {
        color: #0a0a0a !important;
        text-shadow: none !important;
    }
    
    .stMetric .stMetricValue {
        color: #0a0a0a !important;
        text-shadow: none !important;
    }

    /* ============================================================
       DIVIDERS - Subtle
       ============================================================ */
    
    hr {
        border: none !important;
        height: 1px !important;
        background: linear-gradient(90deg, 
            transparent 0%, 
            rgba(0, 0, 0, 0.08) 30%, 
            rgba(0, 0, 0, 0.12) 50%, 
            rgba(0, 0, 0, 0.08) 70%, 
            transparent 100%
        ) !important;
        margin: 1.5rem 0 !important;
    }

    /* ============================================================
       SCROLLBAR - Dark
       ============================================================ */
    
    ::-webkit-scrollbar {
        width: 8px;
        height: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: rgba(30, 30, 30, 0.4) !important;
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: rgba(80, 80, 80, 0.5) !important;
        border-radius: 4px;
        border: 1px solid rgba(0, 0, 0, 0.1);
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: rgba(80, 80, 80, 0.7) !important;
    }

    /* ============================================================
       ALERT MESSAGES - Apple Liquid Glass style
       ============================================================ */
    
    .stAlert {
        background: rgba(255, 255, 255, 0.06) !important;
        backdrop-filter: blur(12px) !important;
        -webkit-backdrop-filter: blur(12px) !important;
        border: 1px solid rgba(255, 255, 255, 0.08) !important;
        border-radius: 0.75rem !important;
        color: #0a0a0a !important;
        text-shadow: none !important;
    }
    
    .stAlert .stMarkdown,
    .stAlert div, 
    .stAlert p, 
    .stAlert span {
        color: #0a0a0a !important;
        text-shadow: none !important;
    }

    /* ============================================================
       LINKS - Deep Black
       ============================================================ */
    
    .stMarkdown a {
        color: #0a0a0a !important;
        text-decoration: underline;
        text-underline-offset: 2px;
        text-shadow: none !important;
    }
    
    .stMarkdown a:hover {
        color: #333333 !important;
    }

    /* ============================================================
       BLOCKQUOTE / STREAMLIT DEFAULT OVERRIDES - DARK
       ============================================================ */
    
    .st-emotion-cache-1r6slb0 {
        background: rgba(0,0,0,0.95) !important;
    }
    
    [data-testid="stHeader"] {
        background: rgba(0,0,0,0.95) !important;
    }
    
    .st-emotion-cache-12fmjuu {
        background: rgba(0,0,0,0.95) !important;
    }
</style>
""", unsafe_allow_html=True)


# ============================================================
# PURE PYTHON BM25 IMPLEMENTATION
# ============================================================
class PureBM25:
    def __init__(self, corpus: List[List[str]], k1: float = 1.5, b: float = 0.75):
        self.k1 = k1
        self.b = b
        self.corpus = corpus
        self.doc_count = len(corpus)
        self.doc_lengths = [len(doc) for doc in corpus]
        self.avg_doc_length = sum(self.doc_lengths) / self.doc_count if self.doc_count > 0 else 0
        
        doc_freq = Counter()
        for doc in corpus:
            for term in set(doc):
                doc_freq[term] += 1
        
        self.idf = {}
        for term, freq in doc_freq.items():
            self.idf[term] = math.log((self.doc_count - freq + 0.5) / (freq + 0.5) + 1.0)
        
        self.doc_term_freqs = [Counter(doc) for doc in corpus]
    
    def get_scores(self, query_tokens: List[str]) -> List[float]:
        scores = []
        for doc_idx in range(self.doc_count):
            score = 0.0
            doc_term_freq = self.doc_term_freqs[doc_idx]
            doc_len = self.doc_lengths[doc_idx]
            
            for term in query_tokens:
                if term not in self.idf:
                    continue
                idf = self.idf[term]
                tf = doc_term_freq.get(term, 0)
                if tf == 0:
                    continue
                numerator = tf * (self.k1 + 1)
                denominator = tf + self.k1 * (1 - self.b + self.b * doc_len / self.avg_doc_length)
                score += idf * (numerator / denominator)
            
            scores.append(score)
        
        return scores


# ============================================================
# API KEY / CLIENT
# ============================================================
def get_api_key() -> Optional[str]:
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return st.secrets["GEMINI_API_KEY"]
    except Exception:
        pass
    return os.getenv("GEMINI_API_KEY")


def get_client() -> Optional[genai.Client]:
    if st.session_state.get("client") is not None:
        return st.session_state.client
    api_key = get_api_key()
    if not api_key:
        return None
    try:
        client = genai.Client(api_key=api_key)
        st.session_state.client = client
        return client
    except Exception as e:
        logger.error(f"Failed to initialize Gemini client: {e}")
        return None


# ============================================================
# EMBEDDINGS (Gemini, PyTorch-free)
# ============================================================
class GeminiEmbeddings:
    def __init__(self, client: genai.Client, model: str = EMBED_MODEL):
        self.client = client
        self.model = model

    def _embed_one(self, text: str, task_type: str, max_retries: int = 3) -> List[float]:
        last_err = None
        for attempt in range(max_retries):
            try:
                resp = self.client.models.embed_content(
                    model=self.model,
                    contents=text,
                    config=types.EmbedContentConfig(task_type=task_type),
                )
                return resp.embeddings[0].values
            except Exception as e:
                last_err = e
                logger.warning(f"Embedding attempt {attempt + 1}/{max_retries} failed: {e}")
                time.sleep(1.5 * (attempt + 1))
        raise RuntimeError(f"Embedding failed after {max_retries} attempts: {last_err}")

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return [self._embed_one(t, "RETRIEVAL_DOCUMENT") for t in texts]

    def embed_query(self, text: str) -> List[float]:
        return self._embed_one(text, "RETRIEVAL_QUERY")


# ============================================================
# DOCUMENT LOADING
# ============================================================
def load_pdf(file) -> str:
    reader = pypdf.PdfReader(file)
    pages_text = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception as e:
            logger.warning(f"Failed to extract text from PDF page: {e}")
    return "\n".join(pages_text)


def load_docx(file) -> str:
    document = docx.Document(file)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def load_txt(file) -> str:
    raw = file.read()
    if isinstance(raw, bytes):
        return raw.decode("utf-8", errors="ignore")
    return raw


def load_csv(file) -> str:
    df = pd.read_csv(file)
    return df.to_string(index=False)


def parse_uploaded_file(file) -> Tuple[str, str]:
    ext = file.name.split(".")[-1].lower()
    if ext == "pdf":
        return load_pdf(file), ext
    elif ext == "docx":
        return load_docx(file), ext
    elif ext == "txt":
        return load_txt(file), ext
    elif ext == "csv":
        return load_csv(file), ext
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


# ============================================================
# CHUNKING
# ============================================================
def chunk_document(text: str, filename: str, filetype: str) -> List[Dict]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE_TOKENS * CHARS_PER_TOKEN,
        chunk_overlap=CHUNK_OVERLAP_TOKENS * CHARS_PER_TOKEN,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    raw_chunks = splitter.split_text(text)
    docs = []
    for i, chunk in enumerate(raw_chunks):
        if not chunk.strip():
            continue
        docs.append({
            "id": f"{filename}::chunk_{i}",
            "text": chunk.strip(),
            "metadata": {
                "filename": filename,
                "filetype": filetype,
                "chunk_index": i,
            },
        })
    return docs


# ============================================================
# HYBRID VECTOR STORE
# ============================================================
class HybridVectorStore:
    def __init__(self, embeddings: GeminiEmbeddings):
        self.embeddings = embeddings
        self.nn: Optional[NearestNeighbors] = None
        self.vectors: Optional[np.ndarray] = None
        self.chunks: List[Dict] = []
        self.bm25: Optional[PureBM25] = None

    @staticmethod
    def _tokenize(text: str) -> List[str]:
        return re.findall(r"\w+", text.lower())

    def is_ready(self) -> bool:
        return self.nn is not None and len(self.chunks) > 0

    def add_documents(self, docs: List[Dict]) -> None:
        if not docs:
            return
        texts = [d["text"] for d in docs]
        vectors = self.embeddings.embed_documents(texts)
        arr = np.array(vectors, dtype="float32")
        
        norms = np.linalg.norm(arr, axis=1, keepdims=True)
        arr = arr / norms

        if self.nn is None:
            self.nn = NearestNeighbors(n_neighbors=min(10, len(arr)), metric='cosine')
            self.vectors = arr
        else:
            self.vectors = np.vstack([self.vectors, arr])
        
        self.nn.fit(self.vectors)
        self.chunks.extend(docs)

        tokenized_corpus = [self._tokenize(c["text"]) for c in self.chunks]
        self.bm25 = PureBM25(tokenized_corpus)

    def search(self, query: str, top_k: int = TOP_K_DOCS) -> List[Dict]:
        if not self.is_ready():
            return []

        q_vec = np.array([self.embeddings.embed_query(query)], dtype="float32")
        q_norm = np.linalg.norm(q_vec)
        if q_norm > 0:
            q_vec = q_vec / q_norm
        
        k = min(top_k * 3, len(self.chunks))
        distances, indices = self.nn.kneighbors(q_vec, n_neighbors=k)
        
        sem_map: Dict[int, float] = {}
        for dist, idx in zip(distances[0], indices[0]):
            if idx >= len(self.chunks):
                continue
            similarity = 1.0 - dist
            sem_map[int(idx)] = float(similarity)

        bm25_scores = self.bm25.get_scores(self._tokenize(query))
        max_bm25 = max(bm25_scores) if len(bm25_scores) and max(bm25_scores) > 0 else 1.0
        kw_map = {i: float(s) / max_bm25 for i, s in enumerate(bm25_scores)}

        all_idx = set(sem_map.keys()) | set(kw_map.keys())
        results = []
        for idx in all_idx:
            sem_s = sem_map.get(idx, 0.0)
            kw_s = kw_map.get(idx, 0.0)
            combined = SEMANTIC_WEIGHT * sem_s + KEYWORD_WEIGHT * kw_s
            chunk = self.chunks[idx]
            results.append({
                "text": chunk["text"],
                "metadata": chunk["metadata"],
                "semantic_score": sem_s,
                "keyword_score": kw_s,
                "confidence": combined,
            })

        results.sort(key=lambda r: r["confidence"], reverse=True)
        return results[:top_k]


# ============================================================
# ROUTER AGENT
# ============================================================
@dataclass
class RouteDecision:
    route: str
    confidence: float
    reason: str


def router_agent(query: str, docs_ready: bool) -> RouteDecision:
    q_lower = query.lower()
    matched = [kw for kw in TIME_SENSITIVE_KEYWORDS if kw in q_lower]

    if not docs_ready:
        return RouteDecision(
            route="web",
            confidence=0.90,
            reason="No documents indexed yet; routing straight to web search.",
        )

    if matched:
        confidence = min(0.60 + 0.10 * len(matched), 0.95)
        return RouteDecision(
            route="both",
            confidence=confidence,
            reason=f"Time-sensitive keyword(s) detected: {', '.join(matched)}. "
                   f"Checking both documents and the web.",
        )

    return RouteDecision(
        route="documents",
        confidence=0.75,
        reason="No time-sensitive keywords found; answering from indexed documents.",
    )


def should_use_web_based_on_results(query: str, doc_results: List[Dict], has_documents: bool) -> bool:
    if not has_documents:
        return True
    if not doc_results:
        return True
    
    query_words = set(query.lower().split())
    relevant_results = 0
    
    for result in doc_results[:3]:
        content = result.get('content', '').lower()
        matches = sum(1 for word in query_words if len(word) > 2 and word in content)
        match_ratio = matches / len(query_words) if query_words else 0
        if match_ratio > 0.2:
            relevant_results += 1
    
    if relevant_results < 2:
        return True
    
    web_keywords = ['current', 'today', 'latest', 'news', 'now', 'recent', 'weather', 'temperature']
    if any(keyword in query.lower() for keyword in web_keywords):
        avg_confidence = sum(r.get('confidence', 0) for r in doc_results[:3]) / 3 if doc_results else 0
        if avg_confidence < 0.6:
            return True
    
    return False


# ============================================================
# WEB FALLBACK
# ============================================================
def web_search(query: str, max_results: int = WEB_MAX_RESULTS) -> List[Dict]:
    results = []
    try:
        with DDGS() as ddgs:
            hits = list(ddgs.text(query, max_results=max_results))
        n = max(len(hits), 1)
        for rank, hit in enumerate(hits):
            relevance = max(0.05, 1.0 - (rank / n) * 0.8)
            results.append({
                "title": hit.get("title", "Untitled"),
                "text": hit.get("body", ""),
                "url": hit.get("href", ""),
                "confidence": relevance,
            })
    except Exception as e:
        logger.error(f"Web search failed: {e}")
        st.warning(f"⚠️ Web search failed: {e}")
    return results


# ============================================================
# ANSWER GENERATION
# ============================================================
def build_context(doc_results: List[Dict], web_results: List[Dict]) -> Tuple[str, List[Dict]]:
    sources = []
    context_parts = []
    idx = 1

    for r in doc_results:
        label = f"Source {idx}"
        meta = r["metadata"]
        context_parts.append(
            f"[{label}] (Document: {meta['filename']}, confidence: {r['confidence']:.0%})\n{r['text']}"
        )
        sources.append({
            "label": label,
            "type": "document",
            "confidence": r["confidence"],
            "filename": meta["filename"],
            "text": r["text"][:600],
        })
        idx += 1

    for r in web_results:
        label = f"Source {idx}"
        context_parts.append(
            f"[{label}] (Web: {r['title']} — {r['url']}, confidence: {r['confidence']:.0%})\n{r['text']}"
        )
        sources.append({
            "label": label,
            "type": "web",
            "confidence": r["confidence"],
            "title": r["title"],
            "url": r["url"],
            "text": r["text"][:600],
        })
        idx += 1

    return "\n\n".join(context_parts), sources


def generate_answer(client: genai.Client, query: str, context: str) -> str:
    if not context.strip():
        prompt = (
            f"{SYSTEM_PROMPT}\n\n"
            "No sources were retrieved for this question. Tell the user honestly "
            "that no relevant documents or web results were found, and only answer "
            "from general knowledge if you clearly label it as such (no citations).\n\n"
            f"Question: {query}"
        )
    else:
        prompt = (
            f"{SYSTEM_PROMPT}\n\nSOURCES:\n{context}\n\n"
            f"QUESTION: {query}\n\nANSWER (with [Source X] citations):"
        )

    try:
        resp = client.models.generate_content(model=GEN_MODEL, contents=prompt)
        return resp.text or "⚠️ The model returned an empty response."
    except Exception as e:
        logger.error(f"Answer generation failed: {e}")
        return f"⚠️ Error generating answer: {e}"


# ============================================================
# STREAMLIT UI
# ============================================================
def init_session_state():
    defaults = {
        "chat_history": [],
        "vector_store": None,
        "indexed_files": [],
        "client": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def process_uploaded_files(files, client: genai.Client) -> int:
    if st.session_state.vector_store is None:
        st.session_state.vector_store = HybridVectorStore(GeminiEmbeddings(client))

    all_docs = []
    for f in files:
        if f.name in st.session_state.indexed_files:
            continue
        try:
            text, ext = parse_uploaded_file(f)
            if not text.strip():
                st.warning(f"No extractable text found in {f.name}; skipping.")
                continue
            docs = chunk_document(text, f.name, ext)
            all_docs.extend(docs)
            st.session_state.indexed_files.append(f.name)
        except Exception as e:
            logger.exception(f"Failed to process {f.name}")
            st.error(f"Failed to process {f.name}: {e}")

    if all_docs:
        with st.spinner(f"Embedding {len(all_docs)} chunks with Gemini..."):
            try:
                st.session_state.vector_store.add_documents(all_docs)
            except Exception as e:
                logger.exception("Failed to build index")
                st.error(f"Failed to build index: {e}")
                return 0
    return len(all_docs)


def render_sources(sources: List[Dict]):
    with st.expander(f"📚 Sources ({len(sources)})"):
        for s in sources:
            icon = "📄" if s["type"] == "document" else "🌐"
            header = f"{icon} **{s['label']}** — confidence {s['confidence']:.0%}"
            if s["type"] == "document":
                header += f" — *{s['filename']}*"
            else:
                header += f" — [{s['title']}]({s['url']})"
            st.markdown(header)
            st.caption(s["text"])
            st.markdown("---")


def main():
    init_session_state()

    st.markdown(
        """
        <div class="main-header">
            <span class="gear-icon">⚙️</span>
            <span class="title-text">RAG+WebFallback</span>
        </div>
        """, 
        unsafe_allow_html=True
    )
    st.markdown('<div class="sub-header">Multi-Agent RAG System with Web Fallback • Source Tracking • Confidence Scoring</div>', unsafe_allow_html=True)

    client = get_client()

    with st.sidebar:
        st.header("⚙️ System Status")
        if client:
            st.success("Gemini API connected")
        else:
            st.error("Gemini API key not set")

        vs: Optional[HybridVectorStore] = st.session_state.vector_store
        n_chunks = len(vs.chunks) if vs else 0
        col1, col2 = st.columns(2)
        col1.metric("Indexed chunks", n_chunks)
        col2.metric("Indexed files", len(st.session_state.indexed_files))

        st.divider()
        st.header("📄 Upload Documents")
        uploaded = st.file_uploader(
            "PDF, TXT, DOCX, CSV",
            type=["pdf", "txt", "docx", "csv"],
            accept_multiple_files=True,
        )
        if st.button("Process documents", type="primary", use_container_width=True, disabled=not uploaded):
            if not client:
                st.error("Cannot process documents without a valid Gemini API key.")
            else:
                added = process_uploaded_files(uploaded, client)
                if added:
                    st.success(f"Indexed {added} new chunk(s).")
                else:
                    st.info("No new chunks added.")

        st.divider()
        with st.expander("Advanced settings"):
            st.write(f"Chunk size: ~{CHUNK_SIZE_TOKENS} tokens")
            st.write(f"Semantic weight: {SEMANTIC_WEIGHT} · Keyword weight: {KEYWORD_WEIGHT}")
            c1, c2 = st.columns(2)
            if c1.button("Clear chat", use_container_width=True):
                st.session_state.chat_history = []
                st.rerun()
            if c2.button("Reset KB", use_container_width=True):
                st.session_state.vector_store = None
                st.session_state.indexed_files = []
                st.rerun()

    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg.get("sources"):
                render_sources(msg["sources"])

    query = st.chat_input("Ask a question about your documents or current events...")
    if query:
        if not client:
            st.error("Please configure `GEMINI_API_KEY` before asking questions.")
            st.stop()

        st.session_state.chat_history.append({"role": "user", "content": query})
        with st.chat_message("user"):
            st.markdown(query)

        with st.chat_message("assistant"):
            vs = st.session_state.vector_store
            docs_ready = vs.is_ready() if vs else False
            decision = router_agent(query, docs_ready)
            st.caption(f"🧭 Route: **{decision.route}** ({decision.confidence:.0%} confidence) — {decision.reason}")

            doc_results: List[Dict] = []
            web_results: List[Dict] = []

            if decision.route in ("documents", "both") and docs_ready:
                with st.spinner("Searching documents..."):
                    try:
                        doc_results = vs.search(query, top_k=TOP_K_DOCS)
                    except Exception as e:
                        logger.exception("Document search failed")
                        st.warning(f"Document search failed: {e}")

            if decision.route in ("web", "both"):
                with st.spinner("Searching the web..."):
                    web_results = web_search(query)

            context, sources = build_context(doc_results, web_results)

            with st.spinner("Generating answer with Gemini..."):
                answer = generate_answer(client, query, context)

            st.markdown(answer)
            if sources:
                render_sources(sources)

            st.session_state.chat_history.append({
                "role": "assistant",
                "content": answer,
                "sources": sources,
            })


if __name__ == "__main__":
    main()